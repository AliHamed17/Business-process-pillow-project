from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = Document()
    
    # Title
    title = doc.add_heading('Academic Report: Haifa Municipality Recruitment Process Mining', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Sections
    sections = [
        ("1. Executive Summary", "This report presents a comprehensive process mining analysis of the recruitment and staffing process at the Haifa Municipality. By analyzing over 1.1 million event logs covering a full year, we identified significant bottlenecks in the late-stage approval phases and substantial internal rework within committee-related stages."),
        ("2. Introduction", "The objective of this project is to analyze the 'Job Staffing' (איוש משרה) process within the Haifa Municipality using modern process mining techniques. The goal is to identify sources of delay, understand process variants, and provide operational recommendations for efficiency."),
        ("3. Performance Analysis", """- Average Cycle Time: 18.5 days.
- Max Cycle Time: 365.3 days.
- Top Bottlenecks: Budget Recommendation (Max 121 days), CEO Decision (Max 134 days)."""),
        ("4. Responsible Change Analysis Result", "Cases WITH reassignments: 16.3 days average. Cases WITHOUT: 80.9 days average. Conclusion: Active management through reassignment prevents stagnation."),
        ("5. Operational Recommendations", """1. SLA Implementation for Budgeting (7-day limit).
2. Automated Committee Scheduling.
3. Requirement Validation at Entry to reduce repetitive pings.
4. Escalation Triggers for inactive cases (14 days).
5. Parallelize Salary Simulation with earlier approvals.""")
    ]
    
    for heading, content in sections:
        doc.add_heading(heading, level=1)
        doc.add_paragraph(content)

    # Save
    path = r'c:\Users\ahamed\business process pillow\haifa-municipality-process-mining\docs\academic_report.docx'
    doc.save(path)
    print(f"Report saved to {path}")

if __name__ == "__main__":
    create_report()
